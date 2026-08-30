import io
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from fpdf import FPDF

# ==========================================
# LaTeX Content Extractor / Parser
# ==========================================
def clean_latex_markup(text: str) -> str:
    """Strip LaTeX specific formatting commands into clean readable text."""
    # Replace \textbf{text} -> **text**
    text = re.sub(r'\\textbf\{([^}]+)\}', r'**\1**', text)
    # Replace \textit{text} or \emph{text} -> *text*
    text = re.sub(r'\\(?:textit|emph)\{([^}]+)\}', r'*\1*', text)
    # Replace \underline{text} -> text
    text = re.sub(r'\\underline\{([^}]+)\}', r'\1', text)
    # Replace \href{url}{text} -> text
    text = re.sub(r'\\href\{[^}]+\}\{([^}]+)\}', r'\1', text)
    # Replace \$ -> $
    text = text.replace(r'\$', '$')
    # Replace \% -> %
    text = text.replace(r'\%', '%')
    # Replace \& -> &
    text = text.replace(r'\&', '&')
    # Replace -- with -
    text = text.replace('--', '-')
    # Remove \vspace{...}
    text = re.sub(r'\\vspace\{[^}]+\}', '', text)
    # Remove \small, \large, \Huge, \scshape
    text = re.sub(r'\\(?:small|large|Huge|scshape)\b', '', text)
    # Strip remaining curly braces if standalone
    text = text.replace('{', '').replace('}', '')
    # Strip multiple spaces
    text = re.sub(r'[ \t]+', ' ', text)
    return text.strip()

def parse_latex_resume(latex_code: str):
    """Parse Jake Ryan LaTeX template into a structured Python dictionary."""
    data = {
        "name": "Virat Patel",
        "title": "Product Management",
        "contact_items": [],
        "sections": []
    }
    
    # Extract Name
    name_match = re.search(r'\\textbf\{[^}]*\\scshape\s+([^}]+)\}', latex_code)
    if not name_match:
        name_match = re.search(r'\\textbf\{\\Huge\s*(?:\\scshape)?\s*([^}]+)\}', latex_code)
    if name_match:
        data["name"] = name_match.group(1).strip()

    # Extract Title (e.g. Product Management)
    title_match = re.search(r'\{\\large\s+([^}]+)\}', latex_code)
    if title_match:
        data["title"] = title_match.group(1).strip()
        
    # Extract Contact Bar
    # Looking inside center block
    center_match = re.search(r'\\begin\{center\}(.*?)\\end\{center\}', latex_code, re.DOTALL)
    if center_match:
        center_text = center_match.group(1)
        # Find lines with $|$
        for line in center_text.split("\n"):
            if "$|$" in line or "mailto:" in line or "linkedin" in line:
                raw_items = line.split("$|$")
                for item in raw_items:
                    cleaned = clean_latex_markup(item).replace("$", "").strip()
                    if cleaned and cleaned not in data["contact_items"]:
                        data["contact_items"].append(cleaned)

    # Extract Sections
    section_chunks = re.split(r'\\section\{([^}]+)\}', latex_code)
    # section_chunks[0] is preamble/header, then alternating (section_title, section_body)
    for i in range(1, len(section_chunks), 2):
        sec_title = section_chunks[i].strip()
        sec_body = section_chunks[i+1] if i+1 < len(section_chunks) else ""
        
        section_obj = {
            "title": sec_title.upper(),
            "items": []
        }
        
        # Check for \resumeSubheading
        # \resumeSubheading{arg1}{arg2}{arg3}{arg4}
        subheadings = re.findall(r'\\resumeSubheading\s*\{([^}]*)\}\s*\{([^}]*)\}\s*\{([^}]*)\}\s*\{([^}]*)\}(.*?)(?=\\resumeSubheading|\\resumeProjectHeading|\\end\{itemize\}|\Z)', sec_body, re.DOTALL)
        if subheadings:
            for sh in subheadings:
                c1, c2, r1, r2, body = sh
                bullets = []
                for b in re.findall(r'\\resumeItem\{([^}]+)\}', body):
                    bullets.append(clean_latex_markup(b))
                section_obj["items"].append({
                    "type": "subheading",
                    "top_left": clean_latex_markup(c1),
                    "top_right": clean_latex_markup(c2),
                    "bottom_left": clean_latex_markup(r1),
                    "bottom_right": clean_latex_markup(r2),
                    "bullets": bullets
                })
        
        # Check for \resumeProjectHeading
        # \resumeProjectHeading{\textbf{title} $|$ \emph{tech}}{demo}
        project_headings = re.findall(r'\\resumeProjectHeading\s*\{([^}]*(?:\{[^}]*\}[^}]*)*)\}\s*\{([^}]*(?:\{[^}]*\}[^}]*)*)\}(.*?)(?=\\resumeProjectHeading|\\resumeSubheading|\\end\{itemize\}|\Z)', sec_body, re.DOTALL)
        if project_headings:
            for ph in project_headings:
                left_raw, right_raw, body = ph
                outcome = ""
                outcome_match = re.search(r'\\item\[\]\\small\\textit\{([^}]+)\}', body)
                if outcome_match:
                    outcome = clean_latex_markup(outcome_match.group(1))
                    
                bullets = []
                for b in re.findall(r'\\resumeItem\{([^}]+)\}', body):
                    bullets.append(clean_latex_markup(b))
                    
                section_obj["items"].append({
                    "type": "project",
                    "left": clean_latex_markup(left_raw),
                    "right": clean_latex_markup(right_raw),
                    "outcome": outcome,
                    "bullets": bullets
                })
                
        # Check for Skills lines
        if "Skills" in sec_title or "SKILLS" in sec_title.upper():
            skill_lines = []
            for m in re.findall(r'\\textbf\{([^}]+)\}\s*\{:\s*([^}]+)\}', sec_body):
                category, values = m
                skill_lines.append({
                    "category": clean_latex_markup(category),
                    "values": clean_latex_markup(values)
                })
            if skill_lines:
                section_obj["items"].append({
                    "type": "skills",
                    "skills": skill_lines
                })
                
        # Check for Achievements / standalone bullets
        if not subheadings and not project_headings and not ("Skills" in sec_title):
            bullets = []
            for b in re.findall(r'\\resumeItem\{([^}]+)\}', sec_body):
                bullets.append(clean_latex_markup(b))
            if bullets:
                section_obj["items"].append({
                    "type": "bullets",
                    "bullets": bullets
                })

        data["sections"].append(section_obj)

    return data


# ==========================================
# DOCX Exporter with Jake Ryan Style & 1-Page Layout
# ==========================================
def add_bottom_border(paragraph, color_hex="000000", size_eighth_pt="6"):
    """Add a clean full-width bottom border line in DOCX."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size_eighth_pt))
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_runs(paragraph, text, base_size=Pt(9.5), bold_default=False, italic_default=False, font_name="Calibri"):
    """Tokenize **bold** and *italic* markdown tokens into Word runs."""
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for token in tokens:
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) >= 4:
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.font.size = base_size
            run.font.name = font_name
        elif token.startswith("*") and token.endswith("*") and len(token) >= 2:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
            run.font.size = base_size
            run.font.name = font_name
        else:
            run = paragraph.add_run(token)
            run.bold = bold_default
            run.italic = italic_default
            run.font.size = base_size
            run.font.name = font_name

def latex_to_docx(latex_code: str) -> io.BytesIO:
    """Render Jake Ryan resume from LaTeX into an exact 1-Page Word document."""
    data = parse_latex_resume(latex_code)
    doc = Document()
    
    # 0.4 inch top/bottom and 0.45 inch left/right margins for exact A4 fit
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.45)
        section.right_margin = Inches(0.45)
        
    # 1. Header Name
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(1)
    run_name = p_name.add_run(data["name"].upper())
    run_name.bold = True
    run_name.font.name = "Calibri"
    run_name.font.size = Pt(16)
    
    # 2. Title Subheader
    if data.get("title"):
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after = Pt(2)
        run_title = p_title.add_run(data["title"])
        run_title.font.name = "Calibri"
        run_title.font.size = Pt(11)
        
    # 3. Contact Line
    if data.get("contact_items"):
        p_contact = doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_contact.paragraph_format.space_before = Pt(0)
        p_contact.paragraph_format.space_after = Pt(4)
        contact_str = " | ".join(data["contact_items"])
        run_c = p_contact.add_run(contact_str)
        run_c.font.name = "Calibri"
        run_c.font.size = Pt(8.5)
        run_c.font.color.rgb = RGBColor(71, 85, 105)

    # 4. Render Sections
    for sec in data["sections"]:
        p_sec = doc.add_paragraph()
        p_sec.paragraph_format.space_before = Pt(5)
        p_sec.paragraph_format.space_after = Pt(2)
        p_sec.paragraph_format.keep_with_next = True
        run_sec = p_sec.add_run(sec["title"].upper())
        run_sec.bold = True
        run_sec.font.name = "Calibri"
        run_sec.font.size = Pt(10.5)
        add_bottom_border(p_sec, color_hex="000000", size_eighth_pt="6")
        
        for item in sec["items"]:
            if item["type"] == "subheading":
                # 2-column table for aligned left/right entries
                table = doc.add_table(rows=2, cols=2)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False
                
                # Set column widths (left 5.2 in, right 2.3 in)
                for row in table.rows:
                    row.cells[0].width = Inches(5.2)
                    row.cells[1].width = Inches(2.2)
                    
                # Row 0: Top Left (Bold) & Top Right (Regular)
                cell_00 = table.cell(0, 0).paragraphs[0]
                cell_00.paragraph_format.space_before = Pt(1)
                cell_00.paragraph_format.space_after = Pt(0)
                add_runs(cell_00, item["top_left"], base_size=Pt(9.5), bold_default=True)
                
                cell_01 = table.cell(0, 1).paragraphs[0]
                cell_01.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cell_01.paragraph_format.space_before = Pt(1)
                cell_01.paragraph_format.space_after = Pt(0)
                add_runs(cell_01, item["top_right"], base_size=Pt(9.0))
                
                # Row 1: Bottom Left (Italic) & Bottom Right (Italic)
                cell_10 = table.cell(1, 0).paragraphs[0]
                cell_10.paragraph_format.space_before = Pt(0)
                cell_10.paragraph_format.space_after = Pt(1)
                add_runs(cell_10, item["bottom_left"], base_size=Pt(9.0), italic_default=True)
                
                cell_11 = table.cell(1, 1).paragraphs[0]
                cell_11.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cell_11.paragraph_format.space_before = Pt(0)
                cell_11.paragraph_format.space_after = Pt(1)
                add_runs(cell_11, item["bottom_right"], base_size=Pt(9.0), italic_default=True)
                
                # Bullets
                for b in item["bullets"]:
                    p_b = doc.add_paragraph(style='List Bullet')
                    p_b.paragraph_format.space_before = Pt(0.5)
                    p_b.paragraph_format.space_after = Pt(1)
                    p_b.paragraph_format.line_spacing = 1.05
                    p_b.paragraph_format.left_indent = Inches(0.18)
                    add_runs(p_b, b, base_size=Pt(9.0))

            elif item["type"] == "project":
                # Project row
                table = doc.add_table(rows=1, cols=2)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False
                table.rows[0].cells[0].width = Inches(5.8)
                table.rows[0].cells[1].width = Inches(1.6)
                
                p_left = table.cell(0, 0).paragraphs[0]
                p_left.paragraph_format.space_before = Pt(2)
                p_left.paragraph_format.space_after = Pt(0)
                add_runs(p_left, item["left"], base_size=Pt(9.5))
                
                p_right = table.cell(0, 1).paragraphs[0]
                p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p_right.paragraph_format.space_before = Pt(2)
                p_right.paragraph_format.space_after = Pt(0)
                add_runs(p_right, item["right"], base_size=Pt(9.0), italic_default=True)
                
                if item.get("outcome"):
                    p_out = doc.add_paragraph()
                    p_out.paragraph_format.space_before = Pt(0)
                    p_out.paragraph_format.space_after = Pt(1)
                    p_out.paragraph_format.left_indent = Inches(0.15)
                    add_runs(p_out, item["outcome"], base_size=Pt(8.5), italic_default=True)
                    
                for b in item["bullets"]:
                    p_b = doc.add_paragraph(style='List Bullet')
                    p_b.paragraph_format.space_before = Pt(0.5)
                    p_b.paragraph_format.space_after = Pt(1)
                    p_b.paragraph_format.line_spacing = 1.05
                    p_b.paragraph_format.left_indent = Inches(0.18)
                    add_runs(p_b, b, base_size=Pt(9.0))
                    
            elif item["type"] == "skills":
                for sk in item["skills"]:
                    p_sk = doc.add_paragraph()
                    p_sk.paragraph_format.space_before = Pt(0.5)
                    p_sk.paragraph_format.space_after = Pt(1)
                    p_sk.paragraph_format.line_spacing = 1.05
                    p_sk.paragraph_format.left_indent = Inches(0.12)
                    r_cat = p_sk.add_run(sk["category"] + ": ")
                    r_cat.bold = True
                    r_cat.font.size = Pt(9.0)
                    r_cat.font.name = "Calibri"
                    r_val = p_sk.add_run(sk["values"])
                    r_val.font.size = Pt(9.0)
                    r_val.font.name = "Calibri"
                    
            elif item["type"] == "bullets":
                for b in item["bullets"]:
                    p_b = doc.add_paragraph(style='List Bullet')
                    p_b.paragraph_format.space_before = Pt(0.5)
                    p_b.paragraph_format.space_after = Pt(1)
                    p_b.paragraph_format.line_spacing = 1.05
                    p_b.paragraph_format.left_indent = Inches(0.18)
                    add_runs(p_b, b, base_size=Pt(9.0))

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


# ==========================================
# PDF Exporter with Jake Ryan Style & 1-Page Layout
# ==========================================
class JakeRyanPDF(FPDF):
    def __init__(self):
        super().__init__(orientation="P", unit="mm", format="A4")
        # Exact ~10mm margins for Jake Ryan A4
        self.set_margins(11, 10, 11)
        self.set_auto_page_break(auto=False)

def clean_for_pdf(text: str) -> str:
    """Normalize text characters for standard Helvetica/Arial fonts."""
    replacements = {
        "\u2013": "-", "\u2014": "--", "\u2018": "'", "\u2019": "'",
        "\u201c": '"', "\u201d": '"', "\u2022": "*", "\u2026": "...",
        "\u2122": "TM", "\u00a9": "(c)", "\u00ae": "(r)", "\u00b7": "|",
        "•": "*"
    }
    for orig, rep in replacements.items():
        text = text.replace(orig, rep)
    return text.encode('latin-1', 'replace').decode('latin-1')

def latex_to_pdf(latex_code: str) -> io.BytesIO:
    """Render Jake Ryan resume from LaTeX into an exact 1-Page A4 PDF."""
    data = parse_latex_resume(latex_code)
    pdf = JakeRyanPDF()
    pdf.add_page()
    
    usable_width = pdf.w - 22 # 210 - 22 = 188mm
    
    # 1. Candidate Name (Large Bold Small-Caps)
    pdf.set_font("Helvetica", "B", 15)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(usable_width, 6, clean_for_pdf(data["name"].upper()), align="C", new_x="LMARGIN", new_y="NEXT")
    
    # 2. Product Management Title
    if data.get("title"):
        pdf.set_font("Helvetica", size=10.5)
        pdf.cell(usable_width, 4.5, clean_for_pdf(data["title"]), align="C", new_x="LMARGIN", new_y="NEXT")
        
    # 3. Contact Line
    if data.get("contact_items"):
        pdf.set_font("Helvetica", size=8.5)
        contact_str = " | ".join(data["contact_items"])
        pdf.cell(usable_width, 4.0, clean_for_pdf(contact_str), align="C", new_x="LMARGIN", new_y="NEXT")
        
    # 4. Sections
    line_h = 3.65
    bullet_font_size = 8.8
    
    for sec in data["sections"]:
        pdf.ln(2.2)
        # Section Heading with Small Caps style
        pdf.set_font("Helvetica", "B", 9.8)
        pdf.cell(usable_width, 4.0, clean_for_pdf(sec["title"].upper()), new_x="LMARGIN", new_y="NEXT")
        
        # Horizontal Rule Line (\titlerule)
        rule_y = pdf.get_y()
        pdf.set_draw_color(0, 0, 0)
        pdf.set_line_width(0.3)
        pdf.line(11, rule_y, 11 + usable_width, rule_y)
        pdf.ln(1.2)
        
        for item in sec["items"]:
            if item["type"] == "subheading":
                # Row 1: Company (Bold) & Location (Regular)
                pdf.set_font("Helvetica", "B", 9.2)
                pdf.cell(usable_width * 0.72, line_h + 0.3, clean_for_pdf(item["top_left"]), align="L")
                pdf.set_font("Helvetica", size=8.8)
                pdf.cell(usable_width * 0.28, line_h + 0.3, clean_for_pdf(item["top_right"]), align="R", new_x="LMARGIN", new_y="NEXT")
                
                # Row 2: Role (Italic) & Date Range (Italic)
                if item["bottom_left"] or item["bottom_right"]:
                    pdf.set_font("Helvetica", "I", 8.8)
                    pdf.cell(usable_width * 0.72, line_h, clean_for_pdf(item["bottom_left"]), align="L")
                    pdf.cell(usable_width * 0.28, line_h, clean_for_pdf(item["bottom_right"]), align="R", new_x="LMARGIN", new_y="NEXT")
                    
                # Bullets
                pdf.set_font("Helvetica", size=bullet_font_size)
                for b in item["bullets"]:
                    clean_b = clean_for_pdf(b.replace("**", "").replace("*", ""))
                    pdf.cell(3.8, line_h, "-", align="L")
                    pdf.multi_cell(usable_width - 3.8, line_h, clean_b, new_x="LMARGIN", new_y="NEXT")
                    
            elif item["type"] == "project":
                # Project Heading: Left (Title + Tech) & Right (Demo Link)
                pdf.set_font("Helvetica", "B", 9.2)
                left_clean = clean_for_pdf(item["left"].replace("**", "").replace("*", ""))
                right_clean = clean_for_pdf(item["right"].replace("**", "").replace("*", ""))
                pdf.cell(usable_width * 0.78, line_h + 0.3, left_clean, align="L")
                pdf.set_font("Helvetica", "I", 8.8)
                pdf.cell(usable_width * 0.22, line_h + 0.3, right_clean, align="R", new_x="LMARGIN", new_y="NEXT")
                
                # Business Outcome line
                if item.get("outcome"):
                    pdf.set_font("Helvetica", "I", 8.2)
                    pdf.set_text_color(50, 50, 50)
                    pdf.multi_cell(usable_width, line_h - 0.2, clean_for_pdf(item["outcome"].replace("**", "").replace("*", "")), new_x="LMARGIN", new_y="NEXT")
                    pdf.set_text_color(0, 0, 0)
                    
                # Bullets
                pdf.set_font("Helvetica", size=bullet_font_size)
                for b in item["bullets"]:
                    clean_b = clean_for_pdf(b.replace("**", "").replace("*", ""))
                    pdf.cell(3.8, line_h, "-", align="L")
                    pdf.multi_cell(usable_width - 3.8, line_h, clean_b, new_x="LMARGIN", new_y="NEXT")
                    
            elif item["type"] == "skills":
                for sk in item["skills"]:
                    pdf.set_font("Helvetica", "B", 8.8)
                    cat_text = clean_for_pdf(sk["category"]) + ": "
                    cat_w = pdf.get_string_width(cat_text) + 1.0
                    pdf.cell(cat_w, line_h, cat_text, align="L")
                    pdf.set_font("Helvetica", size=8.8)
                    pdf.multi_cell(usable_width - cat_w, line_h, clean_for_pdf(sk["values"]), new_x="LMARGIN", new_y="NEXT")
                    
            elif item["type"] == "bullets":
                pdf.set_font("Helvetica", size=bullet_font_size)
                for b in item["bullets"]:
                    clean_b = clean_for_pdf(b.replace("**", "").replace("*", ""))
                    pdf.cell(3.8, line_h, "-", align="L")
                    pdf.multi_cell(usable_width - 3.8, line_h, clean_b, new_x="LMARGIN", new_y="NEXT")

    pdf_bytes = pdf.output()
    pdf_io = io.BytesIO(pdf_bytes)
    pdf_io.seek(0)
    return pdf_io


# Backwards compatibility wrappers
def markdown_to_docx(content: str) -> io.BytesIO:
    if "\\documentclass" in content or "\\section" in content:
        return latex_to_docx(content)
    # Convert markdown to latex format or use standard converter
    return latex_to_docx(content)

def markdown_to_pdf(content: str) -> io.BytesIO:
    if "\\documentclass" in content or "\\section" in content:
        return latex_to_pdf(content)
    return latex_to_pdf(content)
